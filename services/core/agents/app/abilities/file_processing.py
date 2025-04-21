from typing import Dict, Any, Optional, BinaryIO
import pandas as pd
import io
import magic
import boto3
from docx import Document
from pypdf import PdfReader
from app.abilities.base import BaseAbility
from app.config import logger, settings
import tabula
from uuid import uuid4
from datetime import datetime
from app.db.models.file import UserFile, FileVersion, FileRelationship
from app.models.session import get_session

class FileProcessingAbility(BaseAbility):
    """Ability for processing and converting different file formats"""
    
    def __init__(self):
        super().__init__()
        self.name = "file_processing"
        self.s3_client = boto3.client('s3')
        self.mime = magic.Magic(mime=True)
        
        # Supported image formats
        self.image_formats = {
            'jpg': 'JPEG',
            'jpeg': 'JPEG',
            'png': 'PNG',
            'gif': 'GIF',
            'bmp': 'BMP',
            'webp': 'WEBP',
            'tiff': 'TIFF'
        }
        
    async def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Execute file processing operations"""
        action = parameters.get("action")
        
        try:
            if action == "convert":
                return await self.convert_file(
                    source_format=parameters["source_format"],
                    target_format=parameters["target_format"],
                    content=parameters["content"],
                    metadata=parameters.get("metadata", {})
                )
            elif action == "extract":
                return await self.extract_content(
                    file_type=parameters["file_type"],
                    content=parameters["content"],
                    extraction_params=parameters.get("extraction_params", {})
                )
            elif action == "store":
                return await self.store_file(
                    content=parameters["content"],
                    filename=parameters["filename"],
                    user_id=parameters.get("user_id"),
                    source=parameters.get("source", "user_upload"),
                    source_task_id=parameters.get("source_task_id"),
                    source_agent_id=parameters.get("source_agent_id"),
                    tags=parameters.get("tags", []),
                    metadata=parameters.get("metadata", {}),
                    related_items=parameters.get("related_items", [])
                )
            elif action == "update":
                return await self.update_file(
                    file_id=parameters["file_id"],
                    updates=parameters["updates"]
                )
            else:
                raise ValueError(f"Unsupported action: {action}")
                
        except Exception as e:
            logger.error(f"Error in file processing: {e}")
            raise

    def _generate_unique_key(self, user_id: str, directory: str, filename: str) -> str:
        """Generate a unique S3 key for a file"""
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        unique_id = str(uuid4())[:8]
        return f"user-content/{user_id}/{directory}/{timestamp}_{unique_id}_{filename}"

    async def store_file(
        self,
        content: bytes,
        filename: str,
        user_id: str,
        source: str = "user_upload",
        source_task_id: Optional[str] = None,
        source_agent_id: Optional[str] = None,
        tags: list = None,
        metadata: Dict[str, Any] = None,
        related_items: list = None
    ) -> Dict[str, Any]:
        """Store file in S3 and track in database"""
        try:
            # Detect MIME type
            mime_type = self.mime.from_buffer(content)
            
            # Generate unique S3 key
            directory = source.lower().replace('_', '-')
            storage_key = self._generate_unique_key(user_id, directory, filename)
            
            # Upload to S3
            buffer = io.BytesIO(content)
            self.s3_client.upload_fileobj(
                buffer,
                settings.DO_STORAGE_BUCKET,
                storage_key,
                ExtraArgs={'ContentType': mime_type}
            )
            
            # Create database record
            async with get_session() as session:
                file_record = UserFile(
                    user_id=user_id,
                    filename=filename,
                    original_filename=filename,
                    storage_key=storage_key,
                    mime_type=mime_type,
                    file_size=len(content),
                    source=source,
                    source_task_id=source_task_id,
                    source_agent_id=source_agent_id,
                    tags=tags or [],
                    metadata=metadata or {}
                )
                session.add(file_record)
                
                # Add relationships if specified
                if related_items:
                    for item in related_items:
                        relationship = FileRelationship(
                            file_id=file_record.id,
                            related_type=item["type"],
                            related_id=item["id"],
                            metadata=item.get("metadata", {})
                        )
                        session.add(relationship)
                
                await session.commit()
                await session.refresh(file_record)
            
            return {
                "id": str(file_record.id),
                "storage_key": storage_key,
                "filename": filename,
                "mime_type": mime_type,
                "size_bytes": len(content),
                "tags": tags or [],
                "metadata": metadata or {}
            }
            
        except Exception as e:
            logger.error(f"Error storing file: {e}")
            raise

    async def update_file(
        self,
        file_id: str,
        updates: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Update file metadata"""
        try:
            async with get_session() as session:
                file_record = await session.get(UserFile, file_id)
                if not file_record:
                    raise ValueError(f"File not found: {file_id}")
                
                # Update allowed fields
                allowed_fields = ["favorite", "tags", "metadata"]
                for field in allowed_fields:
                    if field in updates:
                        setattr(file_record, field, updates[field])
                
                # Update last accessed time if specified
                if updates.get("update_access_time", False):
                    file_record.last_accessed_at = datetime.utcnow()
                
                await session.commit()
                await session.refresh(file_record)
                
                return {
                    "id": str(file_record.id),
                    "storage_key": file_record.storage_key,
                    "filename": file_record.filename,
                    "favorite": file_record.favorite,
                    "tags": file_record.tags,
                    "metadata": file_record.metadata,
                    "updated_at": file_record.updated_at.isoformat()
                }
                
        except Exception as e:
            logger.error(f"Error updating file: {e}")
            raise

    async def convert_file(
        self,
        source_format: str,
        target_format: str,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Convert file from one format to another and track version"""
        try:
            # Perform conversion
            conversion_result = await super().convert_file(
                source_format=source_format,
                target_format=target_format,
                content=content,
                metadata=metadata
            )
            
            # If this is a conversion of an existing file, track the version
            if original_file_id := metadata.get("original_file_id"):
                async with get_session() as session:
                    # Get the next version number
                    latest_version = await session.query(FileVersion)\
                        .filter(FileVersion.original_file_id == original_file_id)\
                        .order_by(FileVersion.version_number.desc())\
                        .first()
                    version_number = (latest_version.version_number + 1) if latest_version else 1
                    
                    # Store the converted file
                    storage_result = await self.store_file(
                        content=conversion_result["content"],
                        filename=f"{metadata.get('filename', 'converted')}.{target_format}",
                        user_id=metadata["user_id"],
                        source="conversion_result",
                        source_task_id=metadata.get("task_id"),
                        source_agent_id=metadata.get("agent_id"),
                        metadata=metadata
                    )
                    
                    # Create version record
                    version = FileVersion(
                        original_file_id=original_file_id,
                        version_number=version_number,
                        storage_key=storage_result["storage_key"],
                        mime_type=conversion_result["mime_type"],
                        file_size=len(conversion_result["content"]),
                        created_by=metadata.get("agent_id", "user"),
                        changes_description=f"Converted from {source_format} to {target_format}",
                        metadata=metadata
                    )
                    session.add(version)
                    await session.commit()
                    
                    conversion_result["version"] = {
                        "number": version_number,
                        "id": str(version.id)
                    }
            
            return conversion_result
            
        except Exception as e:
            logger.error(f"Error in file conversion: {e}")
            raise

    async def extract_content(self, file_type: str, content: bytes, 
                            extraction_params: Dict[str, Any]) -> Dict[str, Any]:
        """Extract structured content from files"""
        buffer = io.BytesIO(content)
        
        try:
            if file_type == "xlsx":
                df = pd.read_excel(buffer)
                return {
                    "data": df.to_dict(orient="records"),
                    "schema": {col: str(dtype) for col, dtype in df.dtypes.items()}
                }
                
            elif file_type == "csv":
                df = pd.read_csv(buffer)
                return {
                    "data": df.to_dict(orient="records"),
                    "schema": {col: str(dtype) for col, dtype in df.dtypes.items()}
                }
                
            elif file_type == "pdf":
                # Extract tables from PDF
                tables = tabula.read_pdf(buffer, pages='all')
                text_content = ""
                
                # Also extract text if requested
                if extraction_params.get("extract_text", True):
                    pdf = PdfReader(buffer)
                    for page in pdf.pages:
                        text_content += page.extract_text() + "\n"
                
                return {
                    "tables": [df.to_dict(orient="records") for df in tables],
                    "text_content": text_content if text_content else None
                }
                
            elif file_type == "docx":
                doc = Document(buffer)
                return {
                    "paragraphs": [p.text for p in doc.paragraphs],
                    "tables": [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]
                }
                
            else:
                raise ValueError(f"Unsupported file type for extraction: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting content: {e}")
            raise
            
    async def get_file(self, storage_key: str) -> BinaryIO:
        """Retrieve file from S3"""
        try:
            buffer = io.BytesIO()
            self.s3_client.download_fileobj(
                settings.DO_STORAGE_BUCKET,
                storage_key,
                buffer
            )
            buffer.seek(0)
            return buffer
        except Exception as e:
            logger.error(f"Error retrieving file: {e}")
            raise
